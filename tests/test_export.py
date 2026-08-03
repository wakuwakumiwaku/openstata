from __future__ import annotations

from pathlib import Path

import pandas as pd
import pytest
from docx import Document
from openpyxl import load_workbook

from openstata import OpenStata, export_table1, table1


@pytest.fixture
def baseline() -> pd.DataFrame:
    patients = pd.DataFrame(
        {
            "arm": ["Control", "Control", "Control", "Treatment", "Treatment", "Treatment"],
            "age": [50, 60, 55, 70, 80, 75],
            "crp": [2.0, 5.0, None, 3.0, 7.0, 10.0],
            "sex": ["Female", "Male", "Female", "Female", "Female", "Male"],
        }
    )
    return table1(
        patients,
        ["age", "crp", "sex"],
        by="arm",
        nonnormal=["crp"],
        pvalues=True,
        standardized_differences=True,
    )


def test_html_export_is_standalone_and_styled(
    baseline: pd.DataFrame, tmp_path: Path
) -> None:
    destination = tmp_path / "table1.html"

    result = export_table1(
        baseline,
        destination,
        title="Table 1. Trial population",
        subtitle="Intention-to-treat cohort",
        footnotes=["Synthetic data only.", "A < B & C."],
    )

    document = destination.read_text(encoding="utf-8")
    assert result == destination.resolve()
    assert "<!doctype html>" in document
    assert "Table 1. Trial population" in document
    assert "Intention-to-treat cohort" in document
    assert "A &lt; B &amp; C." in document
    assert "@media print" in document
    assert "rowspan=\"2\"" in document
    assert ">Control (n=3)<" in document
    assert "arm=Control" not in document
    assert "https://" not in document


def test_export_protects_existing_files(baseline: pd.DataFrame, tmp_path: Path) -> None:
    destination = tmp_path / "table1.html"
    destination.write_text("keep me", encoding="utf-8")

    with pytest.raises(FileExistsError):
        export_table1(baseline, destination)

    export_table1(baseline, destination, overwrite=True)
    assert destination.read_text(encoding="utf-8").startswith("<!doctype html>")


def test_export_rejects_bad_style_and_format(
    baseline: pd.DataFrame, tmp_path: Path
) -> None:
    with pytest.raises(ValueError, match="style"):
        export_table1(baseline, tmp_path / "table.html", style="neon")  # type: ignore[arg-type]
    with pytest.raises(ValueError, match="Unsupported"):
        export_table1(baseline, tmp_path / "table.pdf")


def test_excel_export_has_professional_workbook_features(
    baseline: pd.DataFrame, tmp_path: Path
) -> None:
    destination = tmp_path / "table1.xlsx"

    export_table1(
        baseline,
        destination,
        title="Table 1. Trial population",
        subtitle="Intention-to-treat cohort",
        style="minimal",
    )

    workbook = load_workbook(destination)
    sheet = workbook["Table 1"]
    assert sheet["A1"].value == "Table 1. Trial population"
    assert sheet.freeze_panes == "C5"
    assert sheet.sheet_view.showGridLines is False
    assert any(str(cell_range).startswith("A") for cell_range in sheet.merged_cells.ranges)
    assert any(str(cell_range).startswith("F") for cell_range in sheet.merged_cells.ranges)
    assert sheet["A4"].value == "Variable"
    assert sheet["C4"].value == "Overall (n=6)"
    assert sheet["D4"].value == "Control (n=3)"
    assert sheet.oddFooter.center.text == "OpenStata"


def test_word_export_is_editable_and_structured(
    baseline: pd.DataFrame, tmp_path: Path
) -> None:
    destination = tmp_path / "table1.docx"

    export_table1(
        baseline,
        destination,
        title="Table 1. Trial population",
        subtitle="Intention-to-treat cohort",
        style="journal",
    )

    document = Document(destination)
    assert document.paragraphs[0].text == "Table 1. Trial population"
    assert document.paragraphs[1].text == "Intention-to-treat cohort"
    assert len(document.tables) == 1
    headers = [cell.text for cell in document.tables[0].rows[0].cells]
    assert headers[:4] == [
        "Variable",
        "Level / statistic",
        "Overall (n=6)",
        "Control (n=3)",
    ]
    assert any(paragraph.text == "Notes" for paragraph in document.paragraphs)


def test_wrapper_builds_and_exports_in_one_step(tmp_path: Path) -> None:
    patients = pd.DataFrame(
        {
            "arm": ["A", "A", "B", "B"],
            "age": [50, 60, 55, 65],
            "sex": ["F", "M", "F", "F"],
        }
    )
    destination = tmp_path / "wrapper.html"

    result = OpenStata(patients).export_table1(
        destination,
        ["age", "sex"],
        by="arm",
        pvalues=True,
        title="Participant characteristics",
    )

    assert result == destination.resolve()
    assert "Participant characteristics" in destination.read_text(encoding="utf-8")
