"""Professional baseline table export for clinical reports."""

from __future__ import annotations

from collections.abc import Sequence
from html import escape
from pathlib import Path
from typing import Literal

import pandas as pd

ExportStyle = Literal["clinical", "journal", "minimal"]

_THEMES: dict[str, dict[str, str]] = {
    "clinical": {
        "accent": "0F6B78",
        "dark": "16324F",
        "pale": "EAF5F7",
        "stripe": "F7FAFC",
        "muted": "64748B",
        "border": "CBD5E1",
    },
    "journal": {
        "accent": "111827",
        "dark": "111827",
        "pale": "F3F4F6",
        "stripe": "FAFAFA",
        "muted": "4B5563",
        "border": "9CA3AF",
    },
    "minimal": {
        "accent": "4F46E5",
        "dark": "1E293B",
        "pale": "EEF2FF",
        "stripe": "F8FAFC",
        "muted": "64748B",
        "border": "E2E8F0",
    },
}

_DEFAULT_TITLE = "Table 1. Baseline characteristics"


def _validate_table(table: pd.DataFrame) -> None:
    if not isinstance(table, pd.DataFrame):
        raise TypeError("table must be a pandas DataFrame")
    if not isinstance(table.index, pd.MultiIndex) or table.index.nlevels < 2:
        raise ValueError("export_table1 expects the MultiIndex output returned by table1()")
    if table.empty:
        raise ValueError("Cannot export an empty baseline table")


def _theme(style: str) -> dict[str, str]:
    if style not in _THEMES:
        choices = ", ".join(_THEMES)
        raise ValueError(f"style must be one of: {choices}")
    return _THEMES[style]


def _display(value: object) -> str:
    return "" if pd.isna(value) else str(value)


def _column_heading(table: pd.DataFrame, value: object) -> str:
    text = str(value)
    heading = text.split("=", 1)[1] if "=" in text else text
    group_sizes = table.attrs.get("openstata_group_sizes", {})
    if isinstance(group_sizes, dict) and text in group_sizes:
        return f"{heading} (n={group_sizes[text]})"
    return heading


def _value_rowspans(
    table: pd.DataFrame,
) -> tuple[dict[tuple[int, int], int], set[tuple[int, int]]]:
    spans: dict[tuple[int, int], int] = {}
    skipped: set[tuple[int, int]] = set()
    for start, length, _ in _row_groups(table):
        if length < 2:
            continue
        for column in range(len(table.columns)):
            values = [
                _display(table.iloc[position, column])
                for position in range(start, start + length)
            ]
            if values[0] and all(not value for value in values[1:]):
                spans[(start, column)] = length
                skipped.update((position, column) for position in range(start + 1, start + length))
    return spans, skipped


def _row_groups(table: pd.DataFrame) -> list[tuple[int, int, str]]:
    variables = [_display(index[0]) for index in table.index]
    groups: list[tuple[int, int, str]] = []
    start = 0
    for position in range(1, len(variables) + 1):
        if position == len(variables) or variables[position] != variables[start]:
            groups.append((start, position - start, variables[start]))
            start = position
    return groups


def _export_notes(table: pd.DataFrame, footnotes: Sequence[str] | None) -> list[str]:
    if isinstance(footnotes, str):
        raise TypeError("footnotes must be a sequence of strings, not one string")
    if footnotes is not None:
        return [str(note) for note in footnotes if str(note).strip()]
    notes = ["Values are mean (SD), median [IQR], or n (%), as appropriate."]
    if "P-value" in table.columns:
        notes.append(
            "P-values use the tests documented by OpenStata and are not adjusted for "
            "multiplicity."
        )
    if "SMD" in table.columns:
        notes.append("SMD = absolute standardized mean difference.")
    return notes


def _prepare_destination(destination: str | Path, overwrite: bool) -> Path:
    path = Path(destination).expanduser()
    if path.exists() and not overwrite:
        raise FileExistsError(f"Refusing to overwrite existing file: {path}")
    path.parent.mkdir(parents=True, exist_ok=True)
    return path


def _html_document(
    table: pd.DataFrame,
    *,
    title: str,
    subtitle: str | None,
    footnotes: Sequence[str],
    style: str,
) -> str:
    colors = _theme(style)
    group_starts = {start: (length, variable) for start, length, variable in _row_groups(table)}
    value_spans, skipped_values = _value_rowspans(table)
    body_rows: list[str] = []

    for position, (index, values) in enumerate(table.iterrows()):
        cells: list[str] = []
        group_class = " group-start" if position in group_starts else ""
        stripe_class = " band" if sum(position >= start for start in group_starts) % 2 == 0 else ""
        if position in group_starts:
            length, variable = group_starts[position]
            cells.append(
                f'<th class="variable" scope="rowgroup" rowspan="{length}">'
                f"{escape(variable)}</th>"
            )
        cells.append(f'<th class="level" scope="row">{escape(_display(index[1]))}</th>')
        for column, value in enumerate(values):
            if (position, column) in skipped_values:
                continue
            rowspan = value_spans.get((position, column))
            attribute = f' rowspan="{rowspan}"' if rowspan else ""
            cells.append(f"<td{attribute}>{escape(_display(value))}</td>")
        body_rows.append(f'<tr class="{group_class}{stripe_class}">{"".join(cells)}</tr>')

    headers = "".join(
        ["<th>Variable</th>", "<th>Level / statistic</th>"]
        + [f"<th>{escape(_column_heading(table, column))}</th>" for column in table.columns]
    )
    subtitle_html = f'<p class="subtitle">{escape(subtitle)}</p>' if subtitle else ""
    notes_html = "".join(
        f'<li><span class="note-index">{index}.</span>{escape(note)}</li>'
        for index, note in enumerate(footnotes, start=1)
    )
    notes_section = (
        f'<section class="notes"><h2>Notes</h2><ol>{notes_html}</ol></section>'
        if notes_html
        else ""
    )

    return f"""<!doctype html>
<html lang="en">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <meta name="generator" content="OpenStata">
  <title>{escape(title)}</title>
  <style>
    :root {{
      --accent: #{colors['accent']};
      --dark: #{colors['dark']};
      --pale: #{colors['pale']};
      --stripe: #{colors['stripe']};
      --muted: #{colors['muted']};
      --border: #{colors['border']};
    }}
    * {{ box-sizing: border-box; }}
    body {{
      margin: 0;
      background: #f4f7fa;
      color: #172033;
      font-family: Inter, ui-sans-serif, system-ui, -apple-system, BlinkMacSystemFont,
        "Segoe UI", sans-serif;
      font-size: 14px;
      line-height: 1.45;
    }}
    main {{
      width: min(1400px, calc(100% - 40px));
      margin: 40px auto;
      background: #fff;
      border: 1px solid var(--border);
      border-radius: 14px;
      box-shadow: 0 18px 45px rgba(22, 50, 79, 0.10);
      overflow: hidden;
    }}
    header {{
      padding: 28px 32px 24px;
      border-top: 6px solid var(--accent);
      border-bottom: 1px solid var(--border);
      background: linear-gradient(135deg, #fff 0%, var(--pale) 100%);
    }}
    .eyebrow {{
      margin: 0 0 7px;
      color: var(--accent);
      font-size: 11px;
      font-weight: 800;
      letter-spacing: 0.13em;
      text-transform: uppercase;
    }}
    h1 {{ margin: 0; color: var(--dark); font-size: 27px; line-height: 1.18; }}
    .subtitle {{ margin: 8px 0 0; color: var(--muted); font-size: 14px; }}
    .table-wrap {{ overflow-x: auto; }}
    table {{ width: 100%; border-collapse: separate; border-spacing: 0; }}
    thead th {{
      position: sticky;
      top: 0;
      z-index: 2;
      padding: 12px 14px;
      background: var(--dark);
      color: #fff;
      border-right: 1px solid rgba(255, 255, 255, 0.13);
      font-size: 11px;
      font-weight: 750;
      letter-spacing: 0.04em;
      text-align: right;
      text-transform: uppercase;
      white-space: nowrap;
    }}
    thead th:first-child, thead th:nth-child(2) {{ text-align: left; }}
    thead th:first-child, tbody tr > :first-child {{ padding-left: 32px; }}
    thead th:last-child, tbody tr > :last-child {{ padding-right: 32px; }}
    tbody th, tbody td {{
      padding: 11px 14px;
      border-bottom: 1px solid var(--border);
      color: #273449;
      text-align: right;
      vertical-align: top;
      white-space: nowrap;
    }}
    tbody th {{ text-align: left; }}
    tbody td[rowspan] {{ vertical-align: middle; }}
    tbody tr.band th, tbody tr.band td {{ background: var(--stripe); }}
    tbody tr.group-start th, tbody tr.group-start td {{
      border-top: 1px solid var(--accent);
    }}
    tbody tr:first-child th, tbody tr:first-child td {{ border-top: 0; }}
    th.variable {{
      min-width: 170px;
      color: var(--dark);
      font-weight: 750;
      vertical-align: middle;
      white-space: normal;
    }}
    th.level {{ min-width: 145px; color: var(--muted); font-weight: 520; }}
    .notes {{ padding: 18px 32px 24px; background: #fff; }}
    .notes h2 {{
      margin: 0 0 8px;
      color: var(--dark);
      font-size: 12px;
      letter-spacing: 0.05em;
      text-transform: uppercase;
    }}
    .notes ol {{ margin: 0; padding: 0; list-style: none; color: var(--muted); }}
    .notes li {{ margin: 4px 0; font-size: 12px; }}
    .note-index {{ display: inline-block; width: 21px; color: var(--accent); font-weight: 700; }}
    @media (max-width: 700px) {{
      main {{ width: 100%; margin: 0; border: 0; border-radius: 0; }}
      header, .notes {{ padding-left: 18px; padding-right: 18px; }}
      h1 {{ font-size: 22px; }}
    }}
    @media print {{
      @page {{ size: landscape; margin: 12mm; }}
      body {{ background: #fff; font-size: 10px; }}
      main {{ width: 100%; margin: 0; border: 0; border-radius: 0; box-shadow: none; }}
      header {{ padding: 12px 0; background: #fff; }}
      .table-wrap {{ overflow: visible; }}
      thead th {{ position: static; }}
      tbody th, tbody td {{ padding: 6px 8px; }}
      .notes {{ padding: 12px 0 0; }}
    }}
  </style>
</head>
<body>
  <main>
    <header>
      <p class="eyebrow">Clinical research summary</p>
      <h1>{escape(title)}</h1>
      {subtitle_html}
    </header>
    <div class="table-wrap">
      <table aria-label="{escape(title)}">
        <thead><tr>{headers}</tr></thead>
        <tbody>{''.join(body_rows)}</tbody>
      </table>
    </div>
    {notes_section}
  </main>
</body>
</html>
"""


def _write_html(
    table: pd.DataFrame,
    path: Path,
    *,
    title: str,
    subtitle: str | None,
    footnotes: Sequence[str],
    style: str,
) -> None:
    document = _html_document(
        table,
        title=title,
        subtitle=subtitle,
        footnotes=footnotes,
        style=style,
    )
    path.write_text(document, encoding="utf-8")


def _write_xlsx(
    table: pd.DataFrame,
    path: Path,
    *,
    title: str,
    subtitle: str | None,
    footnotes: Sequence[str],
    style: str,
) -> None:
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
        from openpyxl.utils import get_column_letter
    except ImportError as error:
        raise ImportError(
            "Excel export requires openpyxl. Install it with: pip install openstata[export]"
        ) from error

    colors = _theme(style)
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Table 1"
    sheet.sheet_view.showGridLines = False
    sheet.freeze_panes = "C5" if subtitle else "C4"
    sheet.sheet_properties.tabColor = colors["accent"]

    column_count = len(table.columns) + 2
    sheet.merge_cells(start_row=1, start_column=1, end_row=1, end_column=column_count)
    title_cell = sheet.cell(1, 1, title)
    title_cell.font = Font(name="Aptos Display", size=18, bold=True, color=colors["dark"])
    title_cell.alignment = Alignment(vertical="center")
    sheet.row_dimensions[1].height = 29

    header_row = 3
    if subtitle:
        sheet.merge_cells(start_row=2, start_column=1, end_row=2, end_column=column_count)
        subtitle_cell = sheet.cell(2, 1, subtitle)
        subtitle_cell.font = Font(name="Aptos", size=10, italic=True, color=colors["muted"])
        header_row = 4

    headers = [
        "Variable",
        "Level / statistic",
        *[_column_heading(table, column) for column in table.columns],
    ]
    thin = Side(style="thin", color=colors["border"])
    medium = Side(style="medium", color=colors["accent"])
    for column, value in enumerate(headers, start=1):
        cell = sheet.cell(header_row, column, value)
        cell.fill = PatternFill("solid", fgColor=colors["dark"])
        cell.font = Font(name="Aptos", size=9, bold=True, color="FFFFFF")
        cell.alignment = Alignment(horizontal="left" if column <= 2 else "right")
        cell.border = Border(bottom=medium)
    sheet.row_dimensions[header_row].height = 25

    groups = _row_groups(table)
    group_by_start = {start: (length, variable) for start, length, variable in groups}
    group_number = 0
    for position, (index, values) in enumerate(table.iterrows()):
        row = header_row + 1 + position
        if position in group_by_start:
            group_number += 1
            _, variable = group_by_start[position]
            sheet.cell(row, 1, variable)
        sheet.cell(row, 2, _display(index[1]))
        for offset, value in enumerate(values, start=3):
            sheet.cell(row, offset, _display(value))
        fill_color = colors["stripe"] if group_number % 2 == 0 else "FFFFFF"
        for column in range(1, column_count + 1):
            cell = sheet.cell(row, column)
            cell.fill = PatternFill("solid", fgColor=fill_color)
            cell.font = Font(
                name="Aptos",
                size=9,
                bold=column == 1,
                color=colors["dark"] if column == 1 else "273449",
            )
            cell.alignment = Alignment(
                horizontal="left" if column <= 2 else "right",
                vertical="top",
                wrap_text=column <= 2,
            )
            cell.border = Border(bottom=thin)

    for start, length, variable in groups:
        first_row = header_row + 1 + start
        last_row = first_row + length - 1
        if length > 1:
            sheet.merge_cells(
                start_row=first_row,
                start_column=1,
                end_row=last_row,
                end_column=1,
            )
            merged = sheet.cell(first_row, 1)
            merged.value = variable
            merged.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
        for column in range(1, column_count + 1):
            sheet.cell(first_row, column).border = Border(top=medium, bottom=thin)

    value_spans, _ = _value_rowspans(table)
    for (start, column), length in value_spans.items():
        first_row = header_row + 1 + start
        last_row = first_row + length - 1
        sheet.merge_cells(
            start_row=first_row,
            start_column=column + 3,
            end_row=last_row,
            end_column=column + 3,
        )
        sheet.cell(first_row, column + 3).alignment = Alignment(
            horizontal="right",
            vertical="center",
        )

    widths = [24, 21] + [17] * len(table.columns)
    for column, width in enumerate(widths, start=1):
        sheet.column_dimensions[get_column_letter(column)].width = width

    notes_start = header_row + len(table) + 2
    if footnotes:
        sheet.merge_cells(
            start_row=notes_start,
            start_column=1,
            end_row=notes_start,
            end_column=column_count,
        )
        note_title = sheet.cell(notes_start, 1, "Notes")
        note_title.font = Font(name="Aptos", size=9, bold=True, color=colors["dark"])
        for offset, note in enumerate(footnotes, start=1):
            note_row = notes_start + offset
            sheet.merge_cells(
                start_row=note_row,
                start_column=1,
                end_row=note_row,
                end_column=column_count,
            )
            cell = sheet.cell(note_row, 1, f"{offset}. {note}")
            cell.font = Font(name="Aptos", size=8, color=colors["muted"])
            cell.alignment = Alignment(wrap_text=True, vertical="top")

    sheet.page_setup.orientation = "landscape"
    sheet.page_setup.fitToWidth = 1
    sheet.page_setup.fitToHeight = 0
    sheet.sheet_properties.pageSetUpPr.fitToPage = True
    sheet.print_title_rows = f"1:{header_row}"
    sheet.oddFooter.center.text = "OpenStata"
    sheet.oddFooter.center.size = 8
    sheet.oddFooter.center.color = colors["muted"]
    workbook.save(path)


def _write_docx(
    table: pd.DataFrame,
    path: Path,
    *,
    title: str,
    subtitle: str | None,
    footnotes: Sequence[str],
    style: str,
) -> None:
    try:
        from docx import Document
        from docx.enum.section import WD_ORIENT
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Inches, Pt, RGBColor
    except ImportError as error:
        raise ImportError(
            "Word export requires python-docx. Install it with: pip install openstata[export]"
        ) from error

    colors = _theme(style)

    def rgb(hex_color: str) -> RGBColor:
        return RGBColor.from_string(hex_color)

    def shade(cell: object, color: str) -> None:
        properties = cell._tc.get_or_add_tcPr()  # type: ignore[attr-defined]
        element = OxmlElement("w:shd")
        element.set(qn("w:fill"), color)
        properties.append(element)

    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)

    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(9)

    title_paragraph = document.add_paragraph()
    title_paragraph.paragraph_format.space_after = Pt(3)
    title_run = title_paragraph.add_run(title)
    title_run.bold = True
    title_run.font.name = "Aptos Display"
    title_run.font.size = Pt(18)
    title_run.font.color.rgb = rgb(colors["dark"])
    if subtitle:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(10)
        run = paragraph.add_run(subtitle)
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = rgb(colors["muted"])

    column_count = len(table.columns) + 2
    word_table = document.add_table(rows=len(table) + 1, cols=column_count)
    word_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    word_table.autofit = False
    word_table.style = "Table Grid"

    headers = [
        "Variable",
        "Level / statistic",
        *[_column_heading(table, column) for column in table.columns],
    ]
    header = word_table.rows[0]
    header._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    for column, value in enumerate(headers):
        cell = header.cells[column]
        cell.text = value
        shade(cell, colors["dark"])
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if column <= 1 else WD_ALIGN_PARAGRAPH.RIGHT
        for run in paragraph.runs:
            run.bold = True
            run.font.name = "Aptos"
            run.font.size = Pt(8)
            run.font.color.rgb = rgb("FFFFFF")

    groups = _row_groups(table)
    group_by_start = {start: (length, variable) for start, length, variable in groups}
    group_number = 0
    for position, (index, values) in enumerate(table.iterrows(), start=0):
        row = word_table.rows[position + 1]
        if position in group_by_start:
            group_number += 1
            _, variable = group_by_start[position]
            row.cells[0].text = variable
        row.cells[1].text = _display(index[1])
        for column, value in enumerate(values, start=2):
            row.cells[column].text = _display(value)
        fill_color = colors["stripe"] if group_number % 2 == 0 else "FFFFFF"
        for column, cell in enumerate(row.cells):
            shade(cell, fill_color)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            paragraph = cell.paragraphs[0]
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.LEFT if column <= 1 else WD_ALIGN_PARAGRAPH.RIGHT
            )
            paragraph.paragraph_format.space_after = Pt(0)
            for run in paragraph.runs:
                run.font.name = "Aptos"
                run.font.size = Pt(8.5)
                if column == 0:
                    run.bold = True
                    run.font.color.rgb = rgb(colors["dark"])

    for start, length, variable in groups:
        first_row = start + 1
        if length > 1:
            merged = word_table.cell(first_row, 0).merge(
                word_table.cell(first_row + length - 1, 0)
            )
            merged.text = variable
            merged.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for run in merged.paragraphs[0].runs:
                run.bold = True
                run.font.name = "Aptos"
                run.font.size = Pt(8.5)
                run.font.color.rgb = rgb(colors["dark"])

    value_spans, _ = _value_rowspans(table)
    for (start, column), length in value_spans.items():
        first_row = start + 1
        merged = word_table.cell(first_row, column + 2).merge(
            word_table.cell(first_row + length - 1, column + 2)
        )
        merged.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        merged.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    widths = [Inches(1.65), Inches(1.45)] + [Inches(1.12)] * len(table.columns)
    for row in word_table.rows:
        for column, width in enumerate(widths):
            row.cells[column].width = width

    if footnotes:
        heading = document.add_paragraph()
        heading.paragraph_format.space_before = Pt(9)
        heading.paragraph_format.space_after = Pt(2)
        run = heading.add_run("Notes")
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = rgb(colors["dark"])
        for index, note in enumerate(footnotes, start=1):
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.12)
            paragraph.paragraph_format.first_line_indent = Inches(-0.12)
            paragraph.paragraph_format.space_after = Pt(1)
            run = paragraph.add_run(f"{index}. {note}")
            run.font.size = Pt(8)
            run.font.color.rgb = rgb(colors["muted"])

    core = document.core_properties
    core.title = title
    core.subject = "Clinical baseline characteristics"
    core.author = "OpenStata"
    document.save(path)


def export_table1(
    table: pd.DataFrame,
    destination: str | Path,
    *,
    title: str = _DEFAULT_TITLE,
    subtitle: str | None = None,
    footnotes: Sequence[str] | None = None,
    style: ExportStyle = "clinical",
    overwrite: bool = False,
) -> Path:
    """Export a baseline table as polished HTML, Excel, or Word.

    The output format is inferred from ``destination``. Supported extensions are
    ``.html``, ``.htm``, ``.xlsx``, and ``.docx``. HTML is standalone and
    print-friendly. Excel and Word require the optional ``export`` dependencies.
    """

    _validate_table(table)
    _theme(style)
    path = _prepare_destination(destination, overwrite)
    notes = _export_notes(table, footnotes)
    suffix = path.suffix.lower()

    if suffix in {".html", ".htm"}:
        _write_html(
            table,
            path,
            title=title,
            subtitle=subtitle,
            footnotes=notes,
            style=style,
        )
    elif suffix == ".xlsx":
        _write_xlsx(
            table,
            path,
            title=title,
            subtitle=subtitle,
            footnotes=notes,
            style=style,
        )
    elif suffix == ".docx":
        _write_docx(
            table,
            path,
            title=title,
            subtitle=subtitle,
            footnotes=notes,
            style=style,
        )
    else:
        raise ValueError("Unsupported export format. Use .html, .xlsx, or .docx")

    return path.resolve()
